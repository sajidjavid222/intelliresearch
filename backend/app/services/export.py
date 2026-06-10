"""Export helpers: BibTeX, and proposal -> DOCX/PDF/Markdown.

DOCX/PDF use optional libs (python-docx, reportlab); if absent we fall back to
Markdown so exports always succeed.
"""
from __future__ import annotations

import io
import re

from app.schemas import Paper, Proposal


def _bib_key(p: Paper) -> str:
    first_author = (p.authors[0].split()[-1] if p.authors else "anon").lower()
    first_author = re.sub(r"[^a-z]", "", first_author) or "anon"
    word = re.sub(r"[^a-z]", "", (p.title or "x").lower().split()[0]) if p.title else "x"
    return f"{first_author}{p.year or ''}{word}"


def papers_to_bibtex(papers: list[Paper]) -> str:
    entries = []
    for p in papers:
        key = _bib_key(p)
        authors = " and ".join(p.authors) if p.authors else "Unknown"
        fields = [
            f"  title = {{{p.title}}}",
            f"  author = {{{authors}}}",
        ]
        if p.year:
            fields.append(f"  year = {{{p.year}}}")
        if p.venue:
            fields.append(f"  journal = {{{p.venue}}}")
        if p.doi:
            fields.append(f"  doi = {{{p.doi}}}")
        if p.url:
            fields.append(f"  url = {{{p.url}}}")
        entries.append("@article{" + key + ",\n" + ",\n".join(fields) + "\n}")
    return "\n\n".join(entries)


def papers_to_ris(papers: list[Paper]) -> str:
    """RIS format — importable by Zotero, Mendeley, EndNote, RefWorks."""
    out = []
    for p in papers:
        lines = ["TY  - JOUR", f"TI  - {p.title}"]
        for a in p.authors:
            lines.append(f"AU  - {a}")
        if p.year:
            lines.append(f"PY  - {p.year}")
        if p.venue:
            lines.append(f"JO  - {p.venue}")
        if p.abstract:
            lines.append(f"AB  - {p.abstract}")
        if p.doi:
            lines.append(f"DO  - {p.doi}")
        if p.url:
            lines.append(f"UR  - {p.url}")
        lines.append("ER  - ")
        out.append("\n".join(lines))
    return "\n\n".join(out)


def proposal_to_markdown(p: Proposal) -> str:
    md = [f"# {p.title}\n", "## Problem Statement\n", p.problem_statement, "\n## Objectives\n"]
    md += [f"- {o}" for o in p.objectives]
    md += ["\n## Methodology\n", p.methodology, "\n## Expected Outcomes\n", p.expected_outcomes]
    if p.budget:
        md += ["\n## Budget\n", "| Item | Cost |", "|------|------|"]
        md += [f"| {b.get('item','')} | {b.get('cost','')} |" for b in p.budget]
    return "\n".join(md)


def proposal_to_docx(p: Proposal) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None
    doc = Document()
    doc.add_heading(p.title, level=0)
    doc.add_heading("Problem Statement", level=1)
    doc.add_paragraph(p.problem_statement)
    doc.add_heading("Objectives", level=1)
    for o in p.objectives:
        doc.add_paragraph(o, style="List Bullet")
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(p.methodology)
    doc.add_heading("Expected Outcomes", level=1)
    doc.add_paragraph(p.expected_outcomes)
    if p.budget:
        doc.add_heading("Budget", level=1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text, t.rows[0].cells[1].text = "Item", "Cost"
        for b in p.budget:
            row = t.add_row().cells
            row[0].text, row[1].text = str(b.get("item", "")), str(b.get("cost", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def proposal_to_pdf(p: Proposal) -> bytes | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer,
        )
    except Exception:
        return None
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(p.title, styles["Title"]), Spacer(1, 12)]
    for heading, body in [
        ("Problem Statement", p.problem_statement),
        ("Methodology", p.methodology),
        ("Expected Outcomes", p.expected_outcomes),
    ]:
        story += [Paragraph(heading, styles["Heading1"]),
                  Paragraph(body.replace("\n", "<br/>"), styles["BodyText"]),
                  Spacer(1, 10)]
    story.insert(2, Paragraph("Objectives", styles["Heading1"]))
    story.insert(3, ListFlowable(
        [ListItem(Paragraph(o, styles["BodyText"])) for o in p.objectives],
        bulletType="bullet"))
    doc.build(story)
    return buf.getvalue()
